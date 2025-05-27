import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import math
import re
from io import BytesIO
from zipfile import ZipFile
from datetime import datetime

st.title("发出商品询证函批量生成")

# 上传Excel文件
excel_file = st.file_uploader("上传发出商品Excel文件", type=["xlsx"])
# 上传Word模板文件
template_file = st.file_uploader("上传Word模板文件", type=["docx"])

def clear_table_rows(table):
    for row in table.rows[1:]:
        tbl = row._tr
        tbl.getparent().remove(tbl)

def add_row(table, data_list):
    row = table.add_row()
    for i, val in enumerate(data_list):
        cell = row.cells[i]
        cell.text = str(val)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)

def replace_text_in_paragraphs(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                inline[i].text = inline[i].text.replace(old_text, new_text)

def clean_filename(name):
    return re.sub(r'[\\/:*?"<>|]', '', name)

if excel_file and template_file:
    if st.button("生成询证函文档"):

        # 读取Excel表
        sheet_dates = {
            '2022年发出商品': '2022-12-31',
            '2023年发出商品': '2023-12-31',
            '2024年发出商品': '2024-12-31',
            '2025年3月发出商品': '2025-03-31'
        }

        dfs = []

        for sheet_name, fixed_date in sheet_dates.items():
            try:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                df['截止日期'] = datetime.strptime(fixed_date, '%Y-%m-%d')
                dfs.append(df)
                st.success(f"成功读取表格：{sheet_name}")
            except Exception as e:
                st.warning(f"未能读取表格【{sheet_name}】，已忽略。错误信息：{e}")

        if len(dfs) == 0:
            st.error("没有读取到任何有效表格，请检查文件。")
            st.stop()

        df_all = pd.concat(dfs, ignore_index=True)
        df_all['被询证方'] = df_all['被询证方'].astype(str).str.strip()

        groups = df_all.groupby('被询证方')

        # 把上传的模板文件保存到内存流，方便多次读取
        template_stream = BytesIO(template_file.read())

        # 创建一个内存的zip包
        zip_buffer = BytesIO()
        with ZipFile(zip_buffer, "a") as zip_file:
            for company_name, group_df in groups:
                # 重新从流加载模板，防止覆盖
                template_stream.seek(0)
                doc = Document(template_stream)

                replace_text_in_paragraphs(doc, '往来单位XXX', company_name)

                table = doc.tables[0]
                clear_table_rows(table)

                headers = [cell.text.strip() for cell in table.rows[0].cells]

                for _, row in group_df.iterrows():
                    row_data = []
                    for col_name in headers:
                        if col_name in group_df.columns:
                            val = row[col_name]
                            if pd.isna(val) or (isinstance(val, float) and math.isnan(val)):
                                val = ''
                            else:
                                if hasattr(val, 'strftime'):
                                    val = val.strftime('%Y/%m/%d')
                        else:
                            val = ''
                        row_data.append(val)
                    add_row(table, row_data)

                file_stream = BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)

                filename = f"{clean_filename(company_name)}_发出商品询证函.docx"
                zip_file.writestr(filename, file_stream.read())

        zip_buffer.seek(0)
        st.success("生成完成！请下载压缩包。")
        st.download_button("下载所有询证函.zip", zip_buffer, file_name="询证函文档合集.zip")
else:
    st.info("请上传Excel文件和Word模板文件。")
